"""LangChain document loaders — infrastructure only."""

from __future__ import annotations

import csv
from pathlib import Path

from backend.app.modules.rag.domain.models import ParsedDocument

_PROCESS_POOL_SUFFIXES = {"pdf", "docx"}
_PROCESS_POOL_MIN_BYTES = 512 * 1024


def should_parse_in_process_pool(path: Path) -> bool:
    suffix = path.suffix.lower().lstrip(".")
    if suffix not in _PROCESS_POOL_SUFFIXES:
        return False
    try:
        return path.stat().st_size >= _PROCESS_POOL_MIN_BYTES
    except OSError:
        return False


def load_documents_from_path(file_path: str, metadata: dict | None = None) -> list[ParsedDocument]:
    """Parse a file into ParsedDocument instances using LangChain when available."""
    path = Path(file_path)
    suffix = path.suffix.lower().lstrip(".")
    base_meta = dict(metadata or {})
    base_meta.setdefault("source", str(path))

    if suffix in {"txt", "md"}:
        return [_parse_text_file(path, base_meta)]

    if suffix == "csv":
        return _parse_csv(path, base_meta)

    if suffix == "pdf":
        return _parse_pdf(path, base_meta)

    if suffix == "docx":
        return _parse_docx(path, base_meta)

    raise ValueError(f"Unsupported file type: {suffix}")


def _parse_text_file(path: Path, metadata: dict) -> ParsedDocument:
    try:
        from langchain_community.document_loaders import TextLoader

        docs = TextLoader(str(path), encoding="utf-8").load()
        content = "\n\n".join(doc.page_content for doc in docs)
        merged = {**metadata, **(docs[0].metadata if docs else {})}
        return ParsedDocument(content=content, metadata=merged)
    except ImportError:
        content = path.read_text(encoding="utf-8", errors="replace")
        return ParsedDocument(content=content, metadata=metadata)


def _parse_csv(path: Path, metadata: dict) -> list[ParsedDocument]:
    try:
        from langchain_community.document_loaders import CSVLoader

        docs = CSVLoader(str(path)).load()
        return [
            ParsedDocument(content=doc.page_content, metadata={**metadata, **doc.metadata})
            for doc in docs
        ]
    except ImportError:
        rows: list[ParsedDocument] = []
        with path.open(encoding="utf-8", errors="replace", newline="") as handle:
            reader = csv.reader(handle)
            for idx, row in enumerate(reader):
                rows.append(
                    ParsedDocument(
                        content=", ".join(row),
                        metadata={**metadata, "row": idx},
                    )
                )
        return rows


def _parse_pdf(path: Path, metadata: dict) -> list[ParsedDocument]:
    try:
        from langchain_community.document_loaders import PyPDFLoader

        docs = PyPDFLoader(str(path)).load()
        return [
            ParsedDocument(
                content=doc.page_content,
                metadata={**metadata, **doc.metadata},
            )
            for doc in docs
            if doc.page_content.strip()
        ]
    except ImportError:
        try:
            from pypdf import PdfReader

            reader = PdfReader(str(path))
            pages: list[ParsedDocument] = []
            for idx, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                if text.strip():
                    pages.append(
                        ParsedDocument(
                            content=text,
                            metadata={**metadata, "page": idx + 1},
                        )
                    )
            return pages
        except ImportError as exc:
            raise ValueError("PDF parsing requires langchain-community or pypdf") from exc


def _parse_docx(path: Path, metadata: dict) -> list[ParsedDocument]:
    try:
        from langchain_community.document_loaders import Docx2txtLoader

        docs = Docx2txtLoader(str(path)).load()
        return [
            ParsedDocument(content=doc.page_content, metadata={**metadata, **doc.metadata})
            for doc in docs
            if doc.page_content.strip()
        ]
    except ImportError:
        try:
            import docx

            document = docx.Document(str(path))
            content = "\n".join(p.text for p in document.paragraphs if p.text.strip())
            return [ParsedDocument(content=content, metadata=metadata)]
        except ImportError as exc:
            raise ValueError("DOCX parsing requires langchain-community or python-docx") from exc
